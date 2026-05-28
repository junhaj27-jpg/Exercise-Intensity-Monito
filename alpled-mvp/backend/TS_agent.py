"""
ALPLED TS Agent - Ollama 기반 D10 통합시험 시나리오 생성기

사용:
    python TS_agent.py --model qwen3b --input ../data/sample_requirements.json --ui ../data/sample_ui.json

출력:
    <input_stem>_output.json
    <input_stem>_output.docx
"""
from __future__ import annotations

import argparse
import json
import sys
import time
from pathlib import Path

try:
    from TS_prompt import SYSTEM_PROMPT, build_prompt
except ImportError:
    print("[ERROR] TS_prompt.py 파일이 같은 디렉토리에 있어야 합니다.")
    sys.exit(1)


def get_client(model_type: str, model_name: str | None = None) -> dict:
    try:
        import ollama  # noqa: F401
    except ImportError:
        print("[ERROR] pip install ollama")
        sys.exit(1)

    mapping = {
        "exaone": "exaone3.5:7.8b",
        "exaone2b": "exaone3.5:2.4b",
        "qwen": "qwen2.5:7b",
        "qwen3b": "qwen2.5:3b",
    }
    if model_type not in mapping:
        raise SystemExit(f"지원하지 않는 모델: {model_type}")
    return {"type": "ollama", "model": model_name or mapping[model_type]}


def run_inference(client: dict, messages: list, system_prompt: str) -> str:
    import ollama
    response = ollama.chat(
        model=client["model"],
        messages=[{"role": "system", "content": system_prompt}] + messages,
        format="json",
        options={"temperature": 0, "num_predict": 8192},
    )
    return response["message"]["content"]


def parse_and_validate(raw: str) -> tuple[dict | None, str]:
    text = raw.strip()
    if text.startswith("```"):
        text = "\n".join(text.splitlines()[1:])
    if text.endswith("```"):
        text = text[:-3].strip()
    try:
        data = json.loads(text)
    except Exception as e:
        return None, f"JSON 파싱 실패: {e}"
    if "scenarios" not in data or "cases" not in data:
        return None, "필수 키 누락: scenarios/cases"
    return data, ""


def fill_missing_cases(data: dict) -> dict:
    filled = 0
    cases = data.setdefault("cases", [])
    for scenario in data.get("scenarios", []):
        sid = scenario.get("scenario_id", "")
        sname = scenario.get("scenario_name", "")
        for tc in scenario.get("test_cases", []):
            tcid = tc.get("test_case_id", "")
            procedures = tc.get("test_procedure", [])
            existing = {c.get("sequence"): c for c in cases if c.get("test_case_id") == tcid}
            for idx, proc in enumerate(procedures, start=1):
                if idx not in existing:
                    cases.append({
                        "round": 1,
                        "scenario_id": sid,
                        "scenario_name": sname,
                        "test_case_id": tcid,
                        "sequence": idx,
                        "process_content": proc,
                        "test_item": "(자동 보완 필요)",
                        "precondition": None,
                        "input_data": "(자동 보완 필요)",
                        "expected_result": "(자동 보완 필요)",
                        "screen_id": "",
                        "test_result": None,
                        "note": "자동 보완된 행입니다."
                    })
                    filled += 1
    if filled:
        cases.sort(key=lambda x: (x.get("scenario_id", ""), x.get("test_case_id", ""), x.get("sequence", 0)))
    print(f"[INFO] 자동 보완 cases: {filled}")
    return data


def convert_to_docx(data: dict, output_path: str) -> None:
    try:
        from docx import Document
    except ImportError:
        print("[WARN] python-docx가 없어 DOCX 생성을 건너뜁니다.")
        return
    doc = Document()
    doc.add_heading("통합시험 시나리오", level=1)
    for scenario in data.get("scenarios", []):
        doc.add_heading(f"{scenario.get('scenario_id','')} - {scenario.get('scenario_name','')}", level=2)
        doc.add_paragraph(scenario.get("scenario_description", ""))
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["시험케이스ID", "설명", "절차", "비고"]):
            table.rows[0].cells[i].text = h
        for tc in scenario.get("test_cases", []):
            row = table.add_row().cells
            row[0].text = str(tc.get("test_case_id", ""))
            row[1].text = str(tc.get("test_case_description", ""))
            row[2].text = "\n".join(tc.get("test_procedure", []))
            row[3].text = str(tc.get("note") or "")
        doc.add_paragraph("■ 시험 케이스 상세")
        ct = doc.add_table(rows=1, cols=7)
        ct.style = "Table Grid"
        for i, h in enumerate(["TC ID", "순번", "업무처리", "시험항목", "입력", "예상결과", "화면ID"]):
            ct.rows[0].cells[i].text = h
        for c in data.get("cases", []):
            if c.get("scenario_id") != scenario.get("scenario_id"):
                continue
            row = ct.add_row().cells
            vals = [c.get("test_case_id"), c.get("sequence"), c.get("process_content"), c.get("test_item"), c.get("input_data"), c.get("expected_result"), c.get("screen_id")]
            for i, v in enumerate(vals):
                row[i].text = "" if v is None else str(v)
    doc.save(output_path)
    print(f"[INFO] DOCX 저장 완료: {output_path}")


def load_json(path: str | Path) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def normalize_requirements(data: dict) -> list[dict]:
    if isinstance(data, list):
        return data
    return data.get("requirements") or data.get("requirement") or [data]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--model", default="qwen3b", choices=["exaone", "exaone2b", "qwen", "qwen3b"])
    parser.add_argument("--model-name", default=None)
    parser.add_argument("--input", required=True)
    parser.add_argument("--ui", default=None)
    parser.add_argument("--output", default=None)
    args = parser.parse_args()

    input_path = Path(args.input)
    req_data = load_json(input_path)
    ui_data = load_json(args.ui) if args.ui else None
    requirements = normalize_requirements(req_data)

    client = get_client(args.model, args.model_name)
    all_scenarios, all_cases = [], []
    start = time.time()

    for req in requirements:
        req_id = req.get("requirement_id") or req.get("id") or "REQ"
        print(f"[INFO] 처리 중: {req_id}")
        raw = run_inference(client, build_prompt(req, ui_data), SYSTEM_PROMPT)
        parsed, err = parse_and_validate(raw)
        if err:
            raw_path = f"{input_path.stem}_{req_id}_raw_output.txt"
            Path(raw_path).write_text(raw, encoding="utf-8")
            print(f"[FAIL] {req_id}: {err} / raw 저장: {raw_path}")
            continue
        parsed = fill_missing_cases(parsed)
        all_scenarios.extend(parsed.get("scenarios", []))
        all_cases.extend(parsed.get("cases", []))

    if not all_scenarios:
        print("[ERROR] 생성된 시나리오가 없습니다.")
        sys.exit(1)

    final = {"scenarios": all_scenarios, "cases": all_cases}
    output_path = Path(args.output or f"{input_path.stem}_output.json")
    output_path.write_text(json.dumps(final, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[INFO] JSON 저장 완료: {output_path}")
    convert_to_docx(final, str(output_path.with_suffix(".docx")))
    print(f"[INFO] 전체 처리 완료: {time.time() - start:.1f}초")


if __name__ == "__main__":
    main()
